"""
build_report.py - Generates report.docx with the two required screenshots.

1. Renders pycaret_leaderboard.csv as a styled PNG table.
2. Spins up the FastAPI app, drives /docs with Playwright, executes /predict,
   and screenshots the resulting Swagger UI panel.
3. Assembles report.docx with both images + captions via python-docx.

Run: .venv/Scripts/python.exe build_report.py
"""

import os
import subprocess
import sys
import time
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
import requests
from docx import Document
from docx.shared import Inches, Pt
from playwright.sync_api import sync_playwright

HERE = Path(__file__).parent
LEADERBOARD_CSV = HERE / "pycaret_leaderboard.csv"
LEADERBOARD_PNG = HERE / "report_leaderboard.png"
SWAGGER_PNG = HERE / "report_swagger.png"
REPORT_DOCX = HERE / "report.docx"
PORT = 8765


# -----------------------------------------------------------------------------
# 1. Leaderboard PNG
# -----------------------------------------------------------------------------
def render_leaderboard_png():
    df = pd.read_csv(LEADERBOARD_CSV)
    # Drop TT (Sec) for a cleaner report image; rename for readability
    if "TT (Sec)" in df.columns:
        df = df.drop(columns=["TT (Sec)"])

    fig, ax = plt.subplots(figsize=(13, 4.5))
    ax.axis("off")
    ax.set_title(
        "PyCaret compare_models() leaderboard - Bank Marketing (UCI 222)",
        fontsize=13, weight="bold", pad=14,
    )
    n_cols = len(df.columns)
    # Wider first column for model names
    col_widths = [0.34] + [(0.66 / (n_cols - 1))] * (n_cols - 1)
    table = ax.table(
        cellText=df.round(4).values,
        colLabels=df.columns,
        colWidths=col_widths,
        loc="center",
        cellLoc="center",
    )
    # Left-align the Model column for readability
    for i in range(len(df) + 1):
        cell = table[(i, 0)]
        cell.get_text().set_horizontalalignment("left")
        cell.PAD = 0.04
    table.auto_set_font_size(False)
    table.set_fontsize(10)
    table.scale(1, 1.6)

    # Header styling
    for j, _ in enumerate(df.columns):
        cell = table[(0, j)]
        cell.set_facecolor("#1f3b5b")
        cell.get_text().set_color("white")
        cell.get_text().set_weight("bold")

    # Highlight winner row (row 1 in table coords = first data row)
    for j in range(len(df.columns)):
        cell = table[(1, j)]
        cell.set_facecolor("#d9ead3")
        cell.get_text().set_weight("bold")

    fig.tight_layout()
    fig.savefig(LEADERBOARD_PNG, dpi=180, bbox_inches="tight")
    plt.close(fig)
    print(f"Wrote {LEADERBOARD_PNG.name}")


# -----------------------------------------------------------------------------
# 2. Swagger screenshot
# -----------------------------------------------------------------------------
def start_uvicorn():
    proc = subprocess.Popen(
        [
            sys.executable, "-m", "uvicorn", "main:app",
            "--host", "127.0.0.1", "--port", str(PORT),
            "--log-level", "warning",
        ],
        cwd=str(HERE),
    )
    # Wait for /health
    for _ in range(60):
        try:
            r = requests.get(f"http://127.0.0.1:{PORT}/health", timeout=1)
            if r.ok:
                return proc
        except requests.RequestException:
            pass
        time.sleep(1)
    proc.terminate()
    raise RuntimeError("FastAPI did not become healthy")


def capture_swagger_screenshot():
    proc = start_uvicorn()
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            ctx = browser.new_context(viewport={"width": 1280, "height": 1600})
            page = ctx.new_page()
            page.goto(f"http://127.0.0.1:{PORT}/docs", wait_until="networkidle")

            # Expand POST /predict
            page.click('div[id="operations-default-predict_predict_post"]')
            page.click("button.try-out__btn")
            # Default request body in the Pydantic schema is already populated
            # with our representative client - just execute it.
            page.click("button.execute")
            # Wait for the live "curl" panel that appears AFTER execution
            page.wait_for_selector("div.curl-command", timeout=20000)
            # Wait for the live server response row (the actual returned body)
            page.wait_for_selector(
                "table.live-responses-table tr.response", timeout=20000
            )
            # Loading spinner gone
            page.wait_for_function(
                "() => !document.querySelector('.loading-container')",
                timeout=20000,
            )
            page.wait_for_timeout(1200)  # final paint settle

            # Scroll to top of operation, then screenshot the operation block
            element = page.locator(
                'div[id="operations-default-predict_predict_post"]'
            )
            element.scroll_into_view_if_needed()
            element.screenshot(path=str(SWAGGER_PNG))
            browser.close()
        print(f"Wrote {SWAGGER_PNG.name}")
    finally:
        proc.terminate()
        try:
            proc.wait(timeout=10)
        except subprocess.TimeoutExpired:
            proc.kill()


# -----------------------------------------------------------------------------
# 3. Build report.docx
# -----------------------------------------------------------------------------
def build_docx():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("OIM3641 Assignment 2 - Report", level=0)
    title.alignment = 1  # center

    p = doc.add_paragraph()
    p.add_run("Author: ").bold = True
    p.add_run("Joshua Bell\n")
    p.add_run("Course: ").bold = True
    p.add_run("OIM3641 - AI Driven App Development (Spring 2026)\n")
    p.add_run("Dataset: ").bold = True
    p.add_run("UCI Bank Marketing (id 222), 8,000-row sample\n")
    p.add_run("Repo: ").bold = True
    p.add_run(
        "https://github.com/JovonBell/OIM3641/tree/main/assignments/assignment-02"
    )

    doc.add_heading("1. PyCaret model comparison table", level=1)
    doc.add_paragraph(
        "Output of compare_models() ranked by accuracy. Random Forest "
        "Classifier wins (top row, highlighted) at 0.8998 mean accuracy "
        "across 10-fold CV. The full leaderboard is committed as "
        "pycaret_leaderboard.csv."
    )
    doc.add_picture(str(LEADERBOARD_PNG), width=Inches(6.5))

    doc.add_heading("2. FastAPI Swagger UI - successful test prediction", level=1)
    doc.add_paragraph(
        "Live capture of the Swagger UI at http://127.0.0.1:8765/docs after "
        "clicking 'Try it out' and 'Execute' on POST /predict with the "
        "default request body. Server response: HTTP 200 with the JSON "
        '{"prediction": "no", "score": 0.86}.'
    )
    doc.add_picture(str(SWAGGER_PNG), width=Inches(6.5))

    doc.add_heading("3. Outcome summary", level=1)
    doc.add_paragraph(
        "PyCaret's low-code workflow surfaced Random Forest as the strongest "
        "candidate (0.8998 CV accuracy). Replicating it manually with "
        "scikit-learn (ColumnTransformer + train_test_split + "
        "RandomForestClassifier) on an 80/20 split scored 0.8975 - within "
        "0.2 pts of PyCaret's 10-fold mean, as expected given the "
        "methodological differences. The dataset is heavily imbalanced "
        "(~88% 'no'), so accuracy understates the real challenge: the 'yes' "
        "class recall is only 0.35 in the manual sklearn report. PyCaret "
        "wins for breadth-first model discovery; sklearn wins for an "
        "auditable production pipeline."
    )

    doc.save(REPORT_DOCX)
    print(f"Wrote {REPORT_DOCX.name}")


if __name__ == "__main__":
    render_leaderboard_png()
    capture_swagger_screenshot()
    build_docx()
    print("\nReport build complete.")
